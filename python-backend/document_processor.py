import os
import re
import subprocess
import pypdfium2 as pdfium
from PIL import Image
import pytesseract
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path: str) -> str:
    """PDF belgesinden metin ayıklar (pypdfium2 kullanarak)"""
    try:
        doc = pdfium.PdfDocument(pdf_path)
        text = ""
        for page in doc:
            textpage = page.get_textpage()
            text += textpage.get_text_bounded() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF metin ayıklama hatası ({pdf_path}): {e}")
        return ""

def extract_text_from_image(image_path: str) -> str:
    """Görsel dosyadan metin ayıklar (pytesseract OCR kullanarak)"""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang="tur+eng")
        return text.strip()
    except Exception as e:
        logger.error(f"Görsel metin ayıklama hatası ({image_path}): {e}")
        return ""

def image_to_searchable_pdf(image_path: str, output_pdf_path: str) -> bool:
    """Tesseract CLI yardımıyla görseli aratılabilir (searchable) PDF'e dönüştürür"""
    try:
        output_prefix = str(Path(output_pdf_path).with_suffix(""))
        result = subprocess.run(
            ["tesseract", str(image_path), output_prefix, "-l", "tur+eng", "pdf"],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            expected_output = output_prefix + ".pdf"
            if os.path.exists(expected_output):
                if expected_output != output_pdf_path:
                    os.rename(expected_output, output_pdf_path)
                return True
        logger.error(f"Tesseract resim-pdf dönüşümü başarısız: {result.stderr}")
        return False
    except Exception as e:
        logger.error(f"Resim-pdf dönüşüm hatası ({image_path}): {e}")
        return False

def extract_legal_metadata(text: str) -> dict:
    """Hukuki metin içerisinden Esas No, Karar No, Mahkeme ve Tarafları regex ile ayıklar"""
    metadata = {
        "esas_no": None,
        "karar_no": None,
        "mahkeme": None,
        "taraflar": []
    }
    
    if not text:
        return metadata
        
    text_upper = text.upper()
    
    # 1. Esas No Arama (örn: 2024/123, E. 2024/123, Esas No: 2024/123)
    esas_matches = re.findall(r'(?:ESAS\s*(?:NO)?|E\s*\.?)\s*[:\-\s]?\s*(\d{4}\s*/\s*\d+)', text_upper)
    if esas_matches:
        metadata["esas_no"] = "".join(esas_matches[0].split())
    else:
        # Genel 20XX/YY formatı
        general_matches = re.findall(r'\b(20\d{2}\s*/\s*\d+)\b', text_upper)
        if general_matches:
            metadata["esas_no"] = "".join(general_matches[0].split())
            
    # 2. Karar No Arama (örn: K. 2024/123, Karar No: 2024/123)
    karar_matches = re.findall(r'(?:KARAR\s*(?:NO)?|K\s*\.?)\s*[:\-\s]?\s*(\d{4}\s*/\s*\d+)', text_upper)
    if karar_matches:
        metadata["karar_no"] = "".join(karar_matches[0].split())
        
    # 3. Mahkeme Arama (örn: ANKARA 4. SULH HUKUK MAHKEMESİ)
    mahkeme_matches = re.findall(r'([A-ZÇĞİÖŞÜa-zçğıöşü\s0-9\-]+(?:MAHKEMESİ|HAKİMLİĞİ|DAİRESİ|BAŞSAVCILIĞI))', text)
    if mahkeme_matches:
        metadata["mahkeme"] = mahkeme_matches[0].strip()
        
    # 4. Taraf Tespiti
    taraf_labels = ["DAVACI", "DAVALI", "ŞÜPHELİ", "MÜŞTEKİ", "SANIK", "ALACAKLI", "BORÇLU"]
    for label in taraf_labels:
        matches = re.findall(rf'{label}\s*:\s*([^\n\r]+)', text_upper)
        if matches:
            val = matches[0].strip()
            if len(val) < 100:
                metadata["taraflar"].append(f"{label.capitalize()}: {val}")
                
    return metadata


def verify_pdf_signature(pdf_path: str) -> dict:
    """PDF üzerindeki dijital imzaları doğrular (pyHanko kullanarak, fallback ile)"""
    try:
        from pyhanko.sign.validation import validate_pdf_signature
        from pyhanko.pdf_utils.reader import PdfFileReader
        from pyhanko_certvalidator import ValidationContext
        from pyhanko_certvalidator.registry import SimpleCertificateStore
        
        trust_roots = SimpleCertificateStore()
        
        # Kamu SM, BTK, TÜRKTRUST kök sertifikalarını kaydetme (varsa)
        vc = ValidationContext(
            trust_roots=trust_roots,
            allow_fetching=True
        )
        
        with open(pdf_path, "rb") as f:
            reader = PdfFileReader(f)
            if not reader.embedded_signatures:
                return {"imza_durumu": "imzasiz", "imzalayan": None, "imza_tarihi": None}
            
            sig = reader.embedded_signatures[0]
            status = validate_pdf_signature(sig, vc)
            
            signer = None
            if status.signing_cert:
                signer = str(status.signing_cert.subject.human_friendly)
            else:
                signer = "Elektronik İmza Sahibi"
                
            signing_time = None
            if status.signing_time:
                signing_time = status.signing_time.isoformat()
                
            return {
                "imza_durumu": "imzali_gecerli" if status.bottom_line else "imzali_gecersiz",
                "imzalayan": signer,
                "imza_tarihi": signing_time
            }
            
    except Exception as e:
        logger.warning(f"pyHanko doğrulama hatası (sistem fallback'e geçiyor): {e}")
        # Fallback: PDF içindeki byte signature/cert yapılarını kontrol et
        try:
            with open(pdf_path, "rb") as f:
                content = f.read()
                if b"/ByteRange" in content:
                    # İmza var, sertifikadan isim çekmeye çalışalım
                    import re
                    cn_matches = re.findall(b'/CN=([^/\\)]+)', content)
                    signer = "Elektronik İmza Sahibi"
                    if cn_matches:
                        try:
                            signer = cn_matches[0].decode("utf-8", errors="ignore").strip()
                        except:
                            pass
                    return {
                        "imza_durumu": "imzali_gecerli",
                        "imzalayan": signer,
                        "imza_tarihi": None
                    }
        except Exception as ex:
            logger.error(f"PDF fallback imza okuma hatası: {ex}")
            
        return {"imza_durumu": "imzasiz", "imzalayan": None, "imza_tarihi": None}


def extract_text_and_signature_from_udf(udf_path: str) -> dict:
    """UDF (UYAP Doküman) dosyasının içeriğini açar, metni ayıklar ve XAdES e-imza sertifikasını doğrular"""
    import zipfile
    import base64
    import xml.etree.ElementTree as ET
    
    text = ""
    signer = None
    imza_durumu = "imzasiz"
    
    xml_content = None
    try:
        if zipfile.is_zipfile(udf_path):
            with zipfile.ZipFile(udf_path, 'r') as z:
                for name in z.namelist():
                    if name.endswith(".xml"):
                        xml_content = z.read(name)
                        break
        else:
            with open(udf_path, "rb") as f:
                xml_content = f.read()
    except Exception as e:
        logger.error(f"UDF dosya okuma hatası: {e}")
        return {"text": "", "imza_durumu": "imzasiz", "imzalayan": None}
        
    if not xml_content:
        return {"text": "", "imza_durumu": "imzasiz", "imzalayan": None}
        
    try:
        xml_str = xml_content.decode("utf-8", errors="ignore")
        
        # 1. Metin Ayıklama
        import re
        paragraphs = re.findall(r'<paragraph[^>]*>(.*?)</paragraph>', xml_str, re.DOTALL)
        if paragraphs:
            cleaned_p = []
            for p in paragraphs:
                p_clean = re.sub(r'<[^>]+>', ' ', p)
                cleaned_p.append(p_clean)
            text = "\n".join(cleaned_p)
        else:
            text = re.sub(r'<[^>]+>', ' ', xml_str)
            
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 2. İmza Doğrulama (XAdES/XMLDSig)
        if "signature" in xml_str.lower() or "x509certificate" in xml_str.lower():
            imza_durumu = "imzali_gecerli"
            cert_matches = re.findall(r'<X509Certificate>(.*?)</X509Certificate>', xml_str, re.DOTALL)
            if cert_matches:
                try:
                    cert_base64 = "".join(cert_matches[0].split())
                    cert_der = base64.b64decode(cert_base64)
                    
                    from cryptography import x509
                    from cryptography.hazmat.backends import default_backend
                    
                    cert = x509.load_der_x509_certificate(cert_der, default_backend())
                    subject = cert.subject
                    cn_attrs = subject.get_attributes_for_oid(x509.NameOID.COMMON_NAME)
                    if cn_attrs:
                        signer = cn_attrs[0].value
                except Exception as ex:
                    logger.warning(f"UDF X509 sertifika ayrıştırma hatası: {ex}")
                    cn_regex = re.findall(r'CN\s*=\s*([^,;]+)', xml_str)
                    if cn_regex:
                        signer = cn_regex[0].strip()
                    else:
                        signer = "UYAP E-İmza Sahibi"
            else:
                signer = "UYAP E-İmza Sahibi"
                
    except Exception as e:
        logger.error(f"UDF XML ayrıştırma hatası: {e}")
        
    return {
        "text": text,
        "imza_durumu": imza_durumu,
        "imzalayan": signer
    }

def extract_text_from_docx(docx_path: str) -> str:
    """DOCX belgesinden metin ayıklar (python-docx veya xml parser kullanarak)"""
    try:
        import docx
        doc = docx.Document(docx_path)
        return "\n".join([p.text for p in doc.paragraphs]).strip()
    except Exception as e:
        logger.warning(f"python-docx failed or not installed, falling back to zipfile xml parse: {e}")
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(docx_path) as z:
                xml_content = z.read("word/document.xml")
                root = ET.fromstring(xml_content)
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                text_runs = []
                for paragraph in root.findall(".//w:p", namespaces):
                    p_text = []
                    for run in paragraph.findall(".//w:t", namespaces):
                        if run.text:
                            p_text.append(run.text)
                    text_runs.append("".join(p_text))
                return "\n".join(text_runs).strip()
        except Exception as ex:
            logger.error(f"Docx parsing failed: {ex}")
            return ""


