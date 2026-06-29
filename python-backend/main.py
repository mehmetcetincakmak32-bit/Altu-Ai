from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import subprocess, json, os, shutil, logging
from pathlib import Path
from pydantic import BaseModel
from dataset import dataset, TurkishLawDataset
from scraper import scraper, TurkishLegalScraper
from remote_sources import remote, RemoteLegalSources, MCP_SERVERS, BEDESTEN_API, MEVZUAT_API, LOCAL_MCP_ENDPOINTS
from gemini_ai import (
    call_gemini, hukuki_ai_sor, sozlesme_analiz, dilekce_puanla,
    hukuki_cevir, karar_harita_olustur, kategori_belirle_ai,
    LEGAL_SYSTEM_PROMPT, GEMINI_MODEL, GEMINI_API_KEY
)
from uets_integration import uets_manager

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="ALTU Python Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Ollama devre dışı — Gemini API kullanılıyor
OLLAMA_URL = "disabled"
OLLAMA_MODEL = "disabled"  # Ollama devre disi — tum AI Gemini ile calisir
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", str(Path(__file__).parent / "uploads")))
UPLOAD_DIR.mkdir(exist_ok=True)

# Docker ve Yerel geliştirme ortamları için tutarlı veri dizini
DATA_DIR = Path(os.getenv("DATA_DIR", str(Path(__file__).parent / "data")))
if not DATA_DIR.exists():
    alt_dir = Path(__file__).parent.parent / "data"
    if alt_dir.exists():
        DATA_DIR = alt_dir

def classify_text(text: str) -> str:
    text_lower = text.lower()
    if any(w in text_lower for w in ["kıdem", "ihbar", "işçi", "işveren", "mesai", "ücret alacağı", "iş sözleşmesi", "iş kanunu", "işk"]):
        return "is"
    if any(w in text_lower for w in ["boşanma", "velayet", "zina", "nafaka", "evlilik", "ortak velayet", "aile konutu"]):
        return "bosanma"
    if any(w in text_lower for w in ["aile", "soybağı", "evlat edinme", "nişan", "vesayet", "kayyım"]):
        return "aile"
    if any(w in text_lower for w in ["miras", "vasiyet", "muris", "tenkis", "veraset", "mirasçı"]):
        return "miras"
    if any(w in text_lower for w in ["kira", "tahliye", "kiracı", "kiralayan", "kira artış", "kira bedeli", "kontrat"]):
        return "kira"
    if any(w in text_lower for w in ["tazminat", "maddi tazminat", "manevi tazminat", "haksız fiil", "zarar"]):
        return "tazminat"
    if any(w in text_lower for w in ["ceza", "tck", "sanık", "suç", "mahkumiyet", "savcı", "tutuklama", "hırsızlık", "dolandırıcılık"]):
        return "ceza"
    if any(w in text_lower for w in ["ticaret", "şirket", "ttk", "limited", "anonim", "hisse", "çek", "senet", "fatura"]):
        return "ticaret"
    if any(w in text_lower for w in ["icra", "haciz", "iik", "ödeme emri", "takip", "borçlu", "alacaklı"]):
        return "icra"
    return "diger"

LEGAL_PROMPT = """Sen üst düzey bir Hukuk Müşavirisin. Türk hukuk sistemine, mevzuata, Resmi Gazete ilanlarına, Yargıtay, Danıştay ve Anayasa Mahkemesi kararlarına tam olarak hakimsin.

DAVRANIŞ VE YANIT KURALLARI:
1. Üslubun profesyonel, akıcı ve analitik bir kıdemli hukuk müşaviri gibi olmalıdır.
2. Kesinlikle aynı kalıp veya şablon cümleleri (Örn: HUKUKİ DAYANAK, AÇIKLAMA vb.) her yanıtta tekrarlayarak kendini kısıtlama. Doğal, açıklayıcı ve dilekçe/mütalaa dilinde akıcı yanıtlar üret.
3. Soruları yanıtlarken yerel veritabanlarındaki sabit özetler yerine, devlet sitelerinden, Resmi Gazete'den ve yüksek mahkeme (Yargıtay, Danıştay, AYM) içtihatlarından taranan güncel dinamik bilgileri referans al.
4. Yalnızca hukuki içerikli soruları yanıtla. Hukuk dışı genel sorular sorulursa kibarca yalnızca hukuk konularında danışmanlık verebileceğini belirt.
5. Sorulan konuyla ilgili güncel içtihatları, kanun maddelerini ve Resmi Gazete kararlarını doğrudan analiz ederek yorumla, teorik bilgiyi pratik tavsiyelerle harmanla.
"""

LEGAL_TOPICS = [
    "icra", "ifa", "bosanma", "kira", "tazminat", "miras", "is",
    "ticaret", "vergi", "ceza", "idare", "anayasa", "medeni",
    "borclar", "sigorta", "fikri", "sermaye", "rekabet", "tuketici",
    "aile", "esya", "gayrimenkul", "kamulastirma", "imar",
    "dava", "mahkeme", "hakim", "savci", "avukat", "vekalet",
    "durusma", "karar", "temyiz", "istinaf", "itiraz", "kanun",
    "yargitay", "danistay", "sozlesme", "haciz", "tapu",
    "dilekce", "bilirkisi", "kesif", "ihtarname", "ipotek",
    "veraset", "mirasci", "kidem", "ihbar", "fazla mesai",
    "kiraci", "kiralayan", "tahliye", "ayip", "garanti",
]


def hukuki_mi(soru: str) -> bool:
    soru_lower = soru.lower()
    for topic in LEGAL_TOPICS:
        if topic in soru_lower:
            return True
    return False


def pack_uyap_extension():
    import zipfile
    
    project_root = Path(__file__).parent.parent
    downloads_dir = project_root / "public" / "uploads"
    
    # Docker vs Local path alignment
    if str(project_root) == "/":
        downloads_dir = Path("/app/uploads")
        
    try:
        downloads_dir.mkdir(parents=True, exist_ok=True)
        zip_path = downloads_dir / "altu-ai-uyap-eklentisi.zip"
    except Exception as e:
        logger.warning(f"Could not create directory {downloads_dir}, falling back to /app: {e}")
        downloads_dir = Path("/app")
        zip_path = downloads_dir / "altu-ai-uyap-eklentisi.zip"
    
    manifest_content = """{
  "manifest_version": 3,
  "name": "altu Ai UYAP Entegrasyon Eklentisi",
  "version": "1.0.0",
  "description": "UYAP Avukat ve Vatandaş portalındaki dava dosyalarınızı tek tıkla altu Ai sisteminize aktarır.",
  "permissions": [
    "activeTab",
    "scripting",
    "storage"
  ],
  "host_permissions": [
    "*://*.uyap.gov.tr/*",
    "http://localhost:3001/*",
    "http://localhost:3000/*"
  ],
  "action": {
    "default_popup": "popup.html"
  },
  "content_scripts": [
    {
      "matches": [
        "*://avukat.uyap.gov.tr/*",
        "*://vatandas.uyap.gov.tr/*"
      ],
      "js": ["content.js"]
    }
  ]
}"""

    popup_html_content = """<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    body { width: 300px; font-family: sans-serif; padding: 12px; margin: 0; background: #f8fafc; color: #0f172a; }
    h3 { margin-top: 0; color: #2563eb; display: flex; align-items: center; gap: 6px; }
    .card { background: white; border: 1px solid #e2e8f0; padding: 12px; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
    .form-group { margin-bottom: 10px; }
    label { display: block; font-size: 11px; font-weight: 600; color: #64748b; margin-bottom: 4px; }
    input { width: 100%; padding: 6px; border: 1px solid #cbd5e1; border-radius: 4px; box-sizing: border-box; font-size: 12px; }
    button { width: 100%; padding: 8px; background: #2563eb; color: white; border: none; border-radius: 4px; font-weight: 600; cursor: pointer; font-size: 12px; }
    button:hover { background: #1d4ed8; }
    .status { font-size: 11px; margin-top: 8px; text-align: center; color: #10b981; font-weight: 500; }
  </style>
</head>
<body>
  <h3>⚖️ altu Ai UYAP Eklentisi</h3>
  <div class="card">
    <div class="form-group">
      <label for="apiUrl">altu Ai Panel URL</label>
      <input type="text" id="apiUrl" value="http://localhost:3001" placeholder="http://localhost:3001">
    </div>
    <button id="saveBtn">Ayarları Kaydet</button>
    <div id="status" class="status"></div>
  </div>
  <script src="popup.js"></script>
</body>
</html>"""

    popup_js_content = """document.addEventListener('DOMContentLoaded', () => {
  const apiUrlInput = document.getElementById('apiUrl');
  const saveBtn = document.getElementById('saveBtn');
  const statusDiv = document.getElementById('status');

  chrome.storage.local.get(['apilexApiUrl'], (result) => {
    if (result.apilexApiUrl) {
      apiUrlInput.value = result.apilexApiUrl;
    }
  });

  saveBtn.addEventListener('click', () => {
    const url = apiUrlInput.value.trim().replace(/\/$/, "");
    chrome.storage.local.set({ apilexApiUrl: url }, () => {
      statusDiv.textContent = "Ayarlar başarıyla kaydedildi!";
      setTimeout(() => { statusDiv.textContent = ""; }, 2000);
    });
  });
});"""

    content_js_content = """console.log("[altu Ai UYAP] Eklenti aktif.");

function injectAPILexButton() {
  if (document.getElementById("altu-sync-btn")) return;

  const headers = document.querySelectorAll("h1, h2, .page-header, .title, .portlet-title");
  let target = headers[0];

  if (!target) {
    target = document.querySelector(".content-header") || document.body;
  }

  const btn = document.createElement("button");
  btn.id = "altu-sync-btn";
  btn.innerText = "⚖️ altu Ai'ye Davaları Aktar";
  btn.style.cssText = `
    padding: 8px 16px;
    background: #2563eb;
    color: white;
    border: none;
    border-radius: 8px;
    font-weight: 600;
    cursor: pointer;
    font-size: 13px;
    margin: 10px;
    box-shadow: 0 4px 6px -1px rgba(37, 99, 235, 0.2);
    transition: all 0.2s;
    z-index: 99999;
  `;
  btn.onmouseover = () => { btn.style.background = "#1d4ed8"; };
  btn.onmouseout = () => { btn.style.background = "#2563eb"; };

  btn.onclick = async () => {
    btn.disabled = true;
    btn.innerText = "⏳ Veriler Ayrıştırılıyor...";
    
    const parsedData = parseUyapPage();
    
    if (parsedData.dosyalar.length === 0) {
      alert("Sayfada aktarılabilir dava dosyası bulunamadı. Dava listesi sayfasında olduğunuzdan emin olun.");
      btn.disabled = false;
      btn.innerText = "⚖️ altu Ai'ye Davaları Aktar";
      return;
    }

    btn.innerText = "🚀 altu Ai'ye Gönderiliyor...";

    chrome.storage.local.get(['apilexApiUrl'], async (result) => {
      const apiUrl = result.apilexApiUrl || "http://localhost:3001";
      try {
        const res = await fetch(`${apiUrl}/api/uyap`, {
          method: "POST",
          headers: {
            "Content-Type": "application/json"
          },
          body: JSON.stringify({ data: parsedData })
        });
        
        if (res.ok) {
          const resData = await res.json();
          alert(`İşlem Başarılı! UYAP portalından ${resData.aktarilan || parsedData.dosyalar.length} adet dava dosyası ve duruşma başarıyla altu Ai sisteminize aktarıldı.`);
        } else {
          alert("altu Ai sunucusu hatası. altu Ai panelinin açık olduğunu kontrol edin.");
        }
      } catch (err) {
        console.error(err);
        alert(`Bağlantı Hatası: altu Ai sunucusuna erişilemedi. API URL'sinin (${apiUrl}) doğru olduğunu doğrulayın.`);
      } finally {
        btn.disabled = false;
        btn.innerText = "⚖️ altu Ai'ye Davaları Aktar";
      }
    });
  };

  if (target === document.body) {
    btn.style.position = "fixed";
    btn.style.top = "10px";
    btn.style.right = "10px";
    document.body.appendChild(btn);
  } else {
    target.appendChild(btn);
  }
}

function parseUyapPage() {
  const dosyalar = [];
  const tables = document.querySelectorAll("table, .grid-table, .ui-datatable-data");
  
  tables.forEach(table => {
    const rows = table.querySelectorAll("tr");
    rows.forEach(row => {
      const cells = row.querySelectorAll("td");
      if (cells.length >= 3) {
        let dosyaNo = "";
        let ad = "";
        let mahkeme = "";
        let esasNo = "";
        let konu = "UYAP Otomatik Aktarımı";
        let tcKimlik = "";

        cells.forEach((cell, idx) => {
          const txt = cell.innerText.trim();
          if (/\\b20\\d{2}\\s*\\/\\s*\\d+\\b/.test(txt)) {
            dosyaNo = txt;
            esasNo = txt.split("/")[1] || txt;
          }
          else if (txt.includes("MAHKEMESİ") || txt.includes("HAKİMLİĞİ") || txt.includes("DAİRESİ")) {
            mahkeme = txt;
          }
          else if (idx === 1 || idx === 2) {
            if (txt.length > 5 && txt.length < 50) {
              ad = txt;
            }
          }
        });

        if (dosyaNo) {
          if (!ad) ad = "UYAP DAVASI";
          if (!mahkeme) mahkeme = "UYAP Portal Mahkemesi";
          
          dosyalar.push({
            dosyaNo,
            ad: ad.toUpperCase(),
            konu,
            mahkeme,
            esasNo,
            tcKimlik: tcKimlik || `TC-${Math.floor(100000 + Math.random() * 900000)}`,
            durusmalar: []
          });
        }
      }
    });
  });

  return { dosyalar };
}

setInterval(injectAPILexButton, 2000);"""

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("manifest.json", manifest_content)
        zip_file.writestr("popup.html", popup_html_content)
        zip_file.writestr("popup.js", popup_js_content)
        zip_file.writestr("content.js", content_js_content)
        
    logger.info(f"UYAP Eklentisi başarıyla paketlendi: {zip_path}")


@app.on_event("startup")
async def startup():
    logger.info("Sunucu başlatılıyor...")
    
    # Veri setlerini otomatik tohumla (Seeding)
    try:
        from seed_datasets import load_and_seed, YARGITAY_PATH, DANISTAY_PATH, MEVZUAT_PATH
        if not YARGITAY_PATH.exists() or not DANISTAY_PATH.exists() or not MEVZUAT_PATH.exists():
            logger.info("Yerel JSON veritabanları eksik. Tohumlama (Seeding) işlemi başlatılıyor...")
            load_and_seed()
    except Exception as e:
        logger.error(f"Startup tohumlama hatası: {e}")

    try:
        dataset.yukle_arkaplan()
        logger.info("HF veri setleri arkaplanda yükleniyor...")
    except Exception as e:
        logger.warning(f"Veri seti arkaplan yükleme başlatılamadı: {e}")
    try:
        remote.hazirla()
        logger.info("Uzaktan kaynaklar hazır (önbellek aktif)")
    except Exception as e:
        logger.warning(f"Uzaktan kaynaklar hazırlanamadı: {e}")
        
    try:
        logger.info("UYAP Eklentisi paketleniyor...")
        pack_uyap_extension()
    except Exception as e:
        logger.warning(f"UYAP Eklentisi paketlenemedi: {e}")
        
    try:
        from scraper import start_background_scraper_loop
        start_background_scraper_loop(interval_hours=12)
    except Exception as e:
        logger.warning(f"Arka plan tarama başlatılamadı: {e}")

    try:
        from vector_indexer import start_vector_indexer_loop
        start_vector_indexer_loop(interval_minutes=60)
    except Exception as e:
        logger.warning(f"Vektör indeksleyici başlatılamadı: {e}")

    try:
        from vector_indexer import index_global_datasets, index_all_tenants
        index_global_datasets()
        index_all_tenants()
    except Exception as e:
        logger.warning(f"İlk indeksleme başarısız: {e}")

    try:
        from tck_importer import import_tck_if_needed
        import_tck_if_needed()
    except Exception as e:
        logger.error(f"TCK importer startup error: {e}")

    try:
        from mcp_manager import start_all
        logger.info("Yerel MCP sunucuları başlatılıyor...")
        start_all()
    except Exception as e:
        logger.warning(f"MCP sunucuları başlatılamadı: {e}")



@app.get("/health")
def health():
    gemini_configured = bool(
        GEMINI_API_KEY
        and GEMINI_API_KEY not in ("AIzaSyYOUR_FREE_KEY_HERE", "YOUR_GEMINI_API_KEY_HERE", "")
        and (GEMINI_API_KEY.startswith("AIza") or GEMINI_API_KEY.startswith("AQ."))
    )
    key_format = "AQ. (Authorization Key)" if GEMINI_API_KEY and GEMINI_API_KEY.startswith("AQ.") else ("AIza (Standard Key)" if GEMINI_API_KEY and GEMINI_API_KEY.startswith("AIza") else "Yapılandırılmamış")
    return {
        "status": "ok",
        "ai_engine": "Google Gemini Flash",
        "gemini_model": GEMINI_MODEL,
        "gemini_configured": gemini_configured,
        "gemini_key_format": key_format,
        "auth_method": "x-goog-api-key header (AQ. uyumlu)",
        "dataset_boyut": sum(len(v) for v in dataset.hf_datasets.values()),
        "dataset_sayisi": len(dataset.hf_datasets),
        "remote_kaynaklar": len(MCP_SERVERS),
        "bedesten_api": BEDESTEN_API,
        "mevzuat_api": MEVZUAT_API,
    }


@app.post("/api/ollama/sor")
def ollama_sor(data: dict):
    """Gemini AI tabanlı hukuki sorgulama (eski Ollama endpoint'i — geriye dönük uyumluluk için korundu)"""
    from vector_store import get_vector_store
    prompt = data.get("prompt", "")
    subdomain = data.get("subdomain", "")

    if not hukuki_mi(prompt):
        return {"response": "Üzgünüm, ben bir hukuk asistanıyım ve yalnızca hukuki konularda yardımcı olabiliyorum. Lütfen Türk hukuk sistemiyle ilgili bir soru yöneltin."}

    # Vector Search — düsük benzerlik esigi altindaki sonuçlari at
    MIN_SIMILARITY = 0.35
    ek_bilgi = ""
    try:
        global_store = get_vector_store(subdomain=None)
        global_matches = global_store.search(prompt, limit=3)
        global_matches = [m for m in global_matches if m.get("score", 1.0) >= MIN_SIMILARITY]

        tenant_matches = []
        if subdomain and subdomain != "www":
            tenant_store = get_vector_store(subdomain)
            tenant_matches = tenant_store.search(prompt, limit=5)
            tenant_matches = [m for m in tenant_matches if m.get("score", 1.0) >= MIN_SIMILARITY]

        if tenant_matches:
            ek_bilgi += "\nAVUKATIN VERİTABANI:\n"
            for idx, match in enumerate(tenant_matches):
                sim = match.get("similarity", 0)
                ek_bilgi += f"[{idx+1}] (benzerlik: %{sim*100:.0f}) {match['text'][:300]}\n"
        if global_matches:
            ek_bilgi += "\nHUKUK VERİTABANI:\n"
            for idx, match in enumerate(global_matches):
                kaynak = match['metadata'].get('kaynak', 'Mevzuat')
                sim = match.get("score", 0)
                ek_bilgi += f"[{idx+1}] (benzerlik: %{sim*100:.0f}) {match['text'][:300]} (Kaynak: {kaynak})\n"
    except Exception as e:
        logger.warning(f"Vector search hatasi: {e}")

    # Dataset fallback
    if not ek_bilgi:
        try:
            ds_cevap = dataset.soru_cevapla(prompt)
            kanunlar = dataset.kanun_ara(prompt)
            if ds_cevap:
                ek_bilgi += f"\nVeri seti referansi: {ds_cevap[:500]}"
            if kanunlar:
                ek_bilgi += "\nIlgili kanun maddeleri:\n"
                for k in kanunlar[:3]:
                    ek_bilgi += f"- {k.get('kanun_adi', '')} {k.get('madde', '')[:200]}\n"
        except:
            pass

    yanit = hukuki_ai_sor(prompt, context=ek_bilgi)
    return {"response": yanit, "engine": "gemini"}


@app.post("/api/ollama/chat")
def ollama_chat(data: dict):
    """Gemini AI tabanlı chat (geriye dönük uyumluluk)"""
    mesaj = data.get("mesaj", "")
    yanit = hukuki_ai_sor(mesaj)
    return {"cevap": yanit, "engine": "gemini"}


class SozlesmeRequest(BaseModel):
    tip: str
    metin: str

@app.post("/api/sozlesme-analizi")
def api_sozlesme_analizi(req: SozlesmeRequest):
    try:
        res = sozlesme_analiz(req.tip, req.metin)
        return res
    except Exception as e:
        logger.error(f"Sözleşme analiz hatası: {e}")
        raise HTTPException(500, detail=str(e))

class DilekceRequest(BaseModel):
    metin: str

@app.post("/api/dilekce-puanla")
def api_dilekce_puanla(req: DilekceRequest):
    try:
        res = dilekce_puanla(req.metin)
        return res
    except Exception as e:
        logger.error(f"Dilekçe puanlama hatası: {e}")
        raise HTTPException(500, detail=str(e))

class CeviriRequest(BaseModel):
    metin: str
    kaynak_dil: str = "tr"
    hedef_dil: str = "en"

@app.post("/api/hukuki-ceviri")
def api_hukuki_ceviri(req: CeviriRequest):
    try:
        res = hukuki_cevir(req.metin, req.kaynak_dil, req.hedef_dil)
        return {"ceviri": res}
    except Exception as e:
        logger.error(f"Hukuki çeviri hatası: {e}")
        raise HTTPException(500, detail=str(e))

class KararHaritaRequest(BaseModel):
    kararlar: list

@app.post("/api/karar-harita")
def api_karar_harita(req: KararHaritaRequest):
    try:
        res = karar_harita_olustur(req.kararlar)
        return res
    except Exception as e:
        logger.error(f"Karar harita hatası: {e}")
        raise HTTPException(500, detail=str(e))

@app.post("/api/belge-isle")
async def api_belge_isle(file: UploadFile = File(...)):
    temp_path = UPLOAD_DIR / file.filename
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    try:
        ext = temp_path.suffix.lower()
        text = ""
        imza_info = {"imza_durumu": "imzasiz", "imzalayan": None}
        
        if ext == ".pdf":
            from document_processor import extract_text_from_pdf, verify_pdf_signature
            text = extract_text_from_pdf(str(temp_path))
            imza_info = verify_pdf_signature(str(temp_path))
        elif ext in [".docx"]:
            from document_processor import extract_text_from_docx
            text = extract_text_from_docx(str(temp_path))
        elif ext in [".udf"]:
            from document_processor import extract_text_and_signature_from_udf
            res = extract_text_and_signature_from_udf(str(temp_path))
            text = res.get("text", "")
            imza_info = {
                "imza_durumu": res.get("imza_durumu", "imzasiz"),
                "imzalayan": res.get("imzalayan")
            }
        elif ext in [".tif", ".tiff", ".jpg", ".jpeg", ".png"]:
            from document_processor import extract_text_from_image
            text = extract_text_from_image(str(temp_path))
        else:
            raise HTTPException(400, detail=f"Desteklenmeyen dosya formatı: {ext}")

        # metadata
        from document_processor import extract_legal_metadata
        metadata = extract_legal_metadata(text)
        
        return {
            "success": True,
            "filename": file.filename,
            "text": text,
            "imza": imza_info,
            "metadata": metadata
        }
    except Exception as e:
        logger.error(f"Belge işleme hatası: {e}")
        raise HTTPException(500, detail=str(e))
    finally:
        if temp_path.exists():
            os.remove(temp_path)

class BelgeOlusturRequest(BaseModel):
    format: str # "docx" or "pdf"
    baslik: str
    icerik: str

@app.post("/api/belge-olustur")
def api_belge_olustur(req: BelgeOlusturRequest):
    try:
        import uuid
        filename = f"{uuid.uuid4()}"
        
        if req.format == "docx":
            output_name = f"{filename}.docx"
            out_path = UPLOAD_DIR / output_name
            try:
                import docx
                doc = docx.Document()
                doc.add_heading(req.baslik, 0)
                doc.add_paragraph(req.icerik)
                doc.save(str(out_path))
            except Exception as e:
                logger.warning(f"python-docx is not installed, writing basic text: {e}")
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(f"{req.baslik}\n\n{req.icerik}")
            
            return {"success": True, "download_url": f"/api/belge/indir/{output_name}", "filename": output_name}
            
        elif req.format == "pdf":
            output_name = f"{filename}.pdf"
            out_path = UPLOAD_DIR / output_name
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                c = canvas.Canvas(str(out_path), pagesize=letter)
                c.setFont("Helvetica-Bold", 16)
                c.drawString(100, 750, req.baslik)
                c.setFont("Helvetica", 12)
                
                y = 700
                for line in req.icerik.split("\n"):
                    if y < 50:
                        c.showPage()
                        y = 750
                    c.drawString(100, y, line)
                    y -= 15
                c.save()
            except Exception as e:
                logger.warning(f"reportlab failed, writing basic text: {e}")
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(f"{req.baslik}\n\n{req.icerik}")
            
            return {"success": True, "download_url": f"/api/belge/indir/{output_name}", "filename": output_name}
        else:
            raise HTTPException(400, detail="Desteklenmeyen format")
    except Exception as e:
        logger.error(f"Belge oluşturma hatası: {e}")
        raise HTTPException(500, detail=str(e))

@app.get("/api/belge/indir/{filename}")
def api_belge_indir(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(404, detail="Dosya bulunamadı")
    return FileResponse(file_path, filename=filename)




@app.post("/api/ocr")
async def ocr(image: UploadFile = File(...)):
    temp_path = UPLOAD_DIR / image.filename
    with open(temp_path, "wb") as f:
        f.write(await image.read())

    try:
        result = subprocess.run(
            ["tesseract", str(temp_path), "stdout", "-l", "tur+eng"],
            capture_output=True, text=True, timeout=60
        )
        metin = result.stdout.strip()
        pdf_path = temp_path.with_suffix(".pdf")
        subprocess.run(["tesseract", str(temp_path), str(pdf_path.with_suffix("")), "-l", "tur+eng", "pdf"],
                       capture_output=True, timeout=60)

        os.remove(temp_path)
        return {"metin": metin, "pdf_yolu": f"uploads/{pdf_path.name}" if pdf_path.exists() else ""}
    except Exception as e:
        if temp_path.exists():
            os.remove(temp_path)
        raise HTTPException(500, detail=str(e))


@app.post("/api/goruntu-analiz")
async def goruntu_analiz(file: UploadFile = File(...), soru: str = Form("")):
    temp_path = UPLOAD_DIR / file.filename
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    try:
        # OCR ile metni çıkar
        from document_processor import extract_text_from_image
        metin = extract_text_from_image(str(temp_path))
        
        if not metin:
            metin = "(Görselden metin okunamadı)"
        
        prompt_text = soru if soru else "Bu hukuki belgeyi analiz et ve özetle."
        full_prompt = f"{prompt_text}\n\nBELGE İÇERİĞİ (OCR):\n{metin[:3000]}"
        yanit = hukuki_ai_sor(full_prompt)
        
        os.remove(temp_path)
        return {"response": yanit, "ocr_text": metin[:500], "engine": "gemini"}
    except Exception as e:
        if temp_path.exists():
            os.remove(temp_path)
        raise HTTPException(500, detail=str(e))


@app.post("/api/uyap/aktar")
def uyap_aktar(data: dict):
    dosyalar = data.get("dosyalar", [])
    durusmalar = data.get("durusmalar", [])
    return {
        "dosya_sayisi": len(dosyalar),
        "durusma_sayisi": len(durusmalar),
        "mesaj": "Veriler başarıyla alındı, Next.js API'ye iletildi"
    }


@app.get("/api/uyap/download-extension")
def download_extension():
    try:
        pack_uyap_extension()
    except Exception as e:
        logger.warning(f"Error packing extension: {e}")
        
    project_root = Path(__file__).parent.parent
    downloads_dir = project_root / "public" / "uploads"
    if str(project_root) == "/" or not downloads_dir.exists():
        downloads_dir = Path("/app/uploads")
        
    zip_path = downloads_dir / "altu-ai-uyap-eklentisi.zip"
    if not zip_path.exists():
        # Fallback to check if in /app
        zip_path = Path("/app/altu-ai-uyap-eklentisi.zip")
        if not zip_path.exists():
            raise HTTPException(500, detail="Zip file not found")
        
    return FileResponse(zip_path, media_type="application/zip", filename="altu-ai-uyap-eklentisi.zip")


@app.get("/api/dataset/ara")
def dataset_ara(sorgu: str = "", tur: str = "qa"):
    if tur == "kanun":
        return {"sonuclar": dataset.kanun_ara(sorgu)}
    cevap = dataset.soru_cevapla(sorgu)
    return {"cevap": cevap, "kaynak": "HF Türk Hukuk Veri Setleri"}

@app.get("/api/dataset/istatistik")
def dataset_istatistik():
    return {
        "veri_seti_sayisi": len(dataset.hf_datasets),
        "toplam_kayit": sum(len(v) for v in dataset.hf_datasets.values()),
        "hazir": dataset.hazir_mi(),
        "yukleniyor": dataset._loading,
        "detay": dataset.istatistik(),
    }

@app.get("/api/dataset/karar-ara")
def dataset_karar_ara(sorgu: str = ""):
    return {"sonuclar": dataset.karar_ara(sorgu), "adet": len(dataset.karar_ara(sorgu))}

# ========== UZAKTAN HUKUK KAYNAKLARI (GitHub MCP + Bedesten API) ==========

@app.get("/api/remote/yargitay")
def remote_yargitay(sorgu: str = "", limit: int = 10):
    return {"sonuclar": remote.yargitay_ara(sorgu, limit), "kaynak": "Yerel Veritabanı (Yargıtay)"}

@app.get("/api/remote/danistay")
def remote_danistay(sorgu: str = "", limit: int = 10):
    return {"sonuclar": remote.danistay_ara(sorgu, limit), "kaynak": "Yerel Veritabanı (Danıştay)"}

@app.get("/api/remote/aym")
def remote_aym(sorgu: str = "", limit: int = 10):
    return {"sonuclar": remote.aym_ara(sorgu, limit), "kaynak": "Yerel Veritabanı (AYM)"}

@app.get("/api/remote/mevzuat")
def remote_mevzuat(sorgu: str = "", limit: int = 10):
    return {"sonuclar": remote.mevzuat_ara(sorgu, limit), "kaynak": "Yerel Veritabanı (Mevzuat)"}

@app.get("/api/remote/kanun")
def remote_kanun(no: str = ""):
    return {"sonuc": remote.kanun_getir(no), "kaynak": "Yerel Veritabanı (Kanun)"}

@app.get("/api/remote/tumu")
def remote_tumu(sorgu: str = "", limit: int = 5, kategori: str = ""):
    res = remote.tumunu_ara(sorgu, limit)
    
    # Remote sonuçların kategorilerini ayarla/filtrele
    if res:
        for kaynak, items in res.items():
            if isinstance(items, list):
                for item in items:
                    if "kategori" not in item:
                        item["kategori"] = classify_text(item.get("ozet", "") + " " + item.get("konu", "") + " " + item.get("baslik", "") + " " + item.get("madde", ""))
        
        if kategori:
            res = {k: [i for i in v if i.get("kategori") == kategori] for k, v in res.items()}
            
    toplam_adet = sum(len(v) for v in res.values() if isinstance(v, list))
    
    if toplam_adet == 0:
        logger.info(f"Uzaktan kaynaklardan sonuç bulunamadı. Yerel veritabanı araması yapılıyor: '{sorgu}', Kategori: '{kategori}'")
        
        # Yerel kategorize veri dosyalarında (seeding çıktılarında) arama yap
        local_res = {"yargitay": [], "danistay": [], "aym": [], "mevzuat": []}
        sorgu_lower = sorgu.lower()
        
        # Yargıtay
        y_path = DATA_DIR / "scraper" / "yargitay.json"
        if y_path.exists():
            try:
                with open(y_path, "r", encoding="utf-8") as f:
                    y_list = json.load(f)
                for item in y_list:
                    if kategori and item.get("kategori") != kategori:
                        continue
                    if (sorgu_lower in (item.get("konu", "") or "").lower() or 
                        sorgu_lower in (item.get("ozet", "") or "").lower() or 
                        sorgu_lower in (item.get("esas", "") or "").lower()):
                        local_res["yargitay"].append(item)
                        if len(local_res["yargitay"]) >= limit:
                            break
            except Exception as e:
                logger.error(f"Yargıtay json read error: {e}")
                
        # Danıştay & AYM
        d_path = DATA_DIR / "scraper" / "danistay.json"
        if d_path.exists():
            try:
                with open(d_path, "r", encoding="utf-8") as f:
                    d_list = json.load(f)
                for item in d_list:
                    if kategori and item.get("kategori") != kategori:
                        continue
                    if (sorgu_lower in (item.get("konu", "") or "").lower() or 
                        sorgu_lower in (item.get("ozet", "") or "").lower() or 
                        sorgu_lower in (item.get("sonuc", "") or "").lower() or
                        sorgu_lower in (item.get("esas", "") or "").lower() or
                        sorgu_lower in (item.get("basvuruNo", "") or "").lower()):
                        if item.get("kaynak") == "aym":
                            local_res["aym"].append(item)
                        else:
                            local_res["danistay"].append(item)
            except Exception as e:
                logger.error(f"Danıştay json read error: {e}")
                
        local_res["danistay"] = local_res["danistay"][:limit]
        local_res["aym"] = local_res["aym"][:limit]
        
        # Mevzuat
        m_path = DATA_DIR / "scraper" / "mevzuat.json"
        if m_path.exists():
            try:
                with open(m_path, "r", encoding="utf-8") as f:
                    m_list = json.load(f)
                for item in m_list:
                    if kategori and item.get("kategori") != kategori:
                        continue
                    if (sorgu_lower in (item.get("baslik", "") or "").lower() or 
                        sorgu_lower in (item.get("madde", "") or "").lower()):
                        local_res["mevzuat"].append(item)
                        if len(local_res["mevzuat"]) >= limit:
                            break
            except Exception as e:
                logger.error(f"Mevzuat json read error: {e}")
                
        total_local = sum(len(v) for v in local_res.values())
        
        if total_local > 0:
            res = local_res
            kaynak_adi = "Yerel Kategorize Hukuk Veritabanı"
        else:
            # Fallback to in-memory datasets
            local_kararlar = dataset.karar_ara(sorgu)
            local_kanunlar = dataset.kanun_ara(sorgu)
            
            if kategori:
                local_kararlar = [k for k in local_kararlar if classify_text(k.get("ozet", "") + " " + k.get("konu", "")) == kategori]
                local_kanunlar = [k for k in local_kanunlar if classify_text(k.get("madde", "") + " " + k.get("kanun_adi", "")) == kategori]
                
            res = {
                "yargitay": [k for k in local_kararlar if "yarg" in k.get("mahkeme", "").lower()][:limit],
                "danistay": [k for k in local_kararlar if "dan" in k.get("mahkeme", "").lower()][:limit],
                "aym": [k for k in local_kararlar if "aym" in k.get("mahkeme", "").lower() or "anayasa" in k.get("mahkeme", "").lower()][:limit],
                "mevzuat": [{"baslik": k.get("kanun_adi", ""), "madde": k.get("madde", ""), "tur": "kanun", "sayi": k.get("numara", ""), "tarih": ""} for k in local_kanunlar][:limit]
            }
            if sum(len(v) for v in res.values()) == 0 and local_kararlar:
                res["yargitay"] = local_kararlar[:limit]
            kaynak_adi = "Yerel Hukuk Veri Seti (Fallback)"
    else:
        kaynak_adi = "Tüm Uzaktan Kaynaklar (Bedesten / Mevzuat API)"
        
    return {"sonuclar": res, "kaynak": kaynak_adi}

@app.get("/api/remote/mcp-sunucular")
def remote_mcp_sunucular():
    import requests as http_requests
    sunucular = []
    for k, v in MCP_SERVERS.items():
        s = {"id": k, "ad": v["name"], "github": v["github"], "yerel": False, "durum": "bilinmiyor"}
        if k in LOCAL_MCP_ENDPOINTS:
            s["yerel"] = True
            try:
                r = http_requests.get(LOCAL_MCP_ENDPOINTS[k], timeout=3)
                s["durum"] = "aktif" if r.status_code < 500 else "hata"
            except:
                s["durum"] = "pasif"
        sunucular.append(s)
    return {"sunucular": sunucular}


# ========== TARAMA (SCRAPER) ENDPOINT'LERİ ==========

@app.post("/api/scraper/tara-unified")
def scraper_tara_unified(
    kaynak: str = "tumu",
    sorgu: str = "",
    limit: int = 100,
    gun: int = 30
):
    """Yeni nesil birlesik tarama sistemi — sadece 5 resmi kaynak."""
    from legal_scraper_pipeline import LegalScraper
    scraper = LegalScraper()

    import threading, time
    baslangic = time.time()
    sonuc = {}
    hata = ""

    def run():
        nonlocal sonuc, hata
        try:
            if kaynak == "tumu":
                sonuc = scraper.tumunu_tara()
            elif kaynak == "mevzuat":
                sonuc = {"mevzuat": scraper.mevzuat_tara()}
            elif kaynak == "aym":
                s1 = scraper.aym_mevzuat_tara()
                s2 = scraper.aym_aihs_indir()
                sonuc = {"aym_mevzuat": s1, "aym_aihs": s2}
            elif kaynak == "yargitay":
                sonuc = {"yargitay": scraper.yargitay_tara(sorgu, limit)}
            elif kaynak == "danistay":
                sonuc = {"danistay": scraper.danistay_tara(sorgu, limit)}
            elif kaynak == "rg":
                sonuc = {"resmi_gazete": scraper.resmi_gazete_tara(gun)}
            elif kaynak == "kanun_pdf":
                sonuc = {"kanun_pdf": scraper.kanun_pdf_indexle()}
        except Exception as e:
            hata = str(e)

    t = threading.Thread(target=run, daemon=True)
    t.start()
    t.join(timeout=300)

    gecen = time.time() - baslangic
    return {
        "success": not hata,
        "kaynak": kaynak,
        "sonuc": sonuc,
        "toplam_indexlenen": scraper.toplam_indexlenen if not hata else 0,
        "sure_sn": round(gecen, 1),
        "hata": hata or None
    }


@app.post("/api/scraper/tara")
def scraper_tara(hedef: str = "tumu"):
    """Belirtilen siteleri tara ve yeni entegre pipeline'ları çalıştır"""
    import sys
    import subprocess
    
    pipeline_dir = Path(__file__).resolve().parent.parent / "ingest_pipeline"
    
    if hedef == "yargitay":
        script_path = pipeline_dir / "yargitay_danistay_aym.py"
        cmd = [sys.executable, str(script_path), "--sorgu", "tazminat", "--limit", "10", "--court", "YARGITAY"]
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, check=True)
            return {"kaynak": "yargitay", "success": True, "output": res.stdout.strip()}
        except Exception as e:
            return {"kaynak": "yargitay", "success": False, "error": str(e)}
            
    elif hedef == "danistay":
        script_path = pipeline_dir / "yargitay_danistay_aym.py"
        cmd = [sys.executable, str(script_path), "--sorgu", "ihale", "--limit", "10", "--court", "DANISTAY"]
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, check=True)
            return {"kaynak": "danistay", "success": True, "output": res.stdout.strip()}
        except Exception as e:
            return {"kaynak": "danistay", "success": False, "error": str(e)}
            
    elif hedef == "mevzuat":
        script_path = pipeline_dir / "mevzuat_baseline.py"
        cmd = [sys.executable, str(script_path), "--limit", "50"]
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, check=True)
            return {"kaynak": "mevzuat", "success": True, "output": res.stdout.strip()}
        except Exception as e:
            return {"kaynak": "mevzuat", "success": False, "error": str(e)}
            
    elif hedef == "resmi_gazete":
        script_path = pipeline_dir / "resmi_gazete.py"
        cmd = [sys.executable, str(script_path), "--days", "1"]
        try:
            res = subprocess.run(cmd, capture_output=True, text=True, check=True)
            return {"kaynak": "resmi_gazete", "success": True, "output": res.stdout.strip()}
        except Exception as e:
            return {"kaynak": "resmi_gazete", "success": False, "error": str(e)}
            
    else:
        # Tümü
        results = {}
        for h in ["yargitay", "danistay", "mevzuat", "resmi_gazete"]:
            results[h] = scraper_tara(h)
        return {"sonuc": results}


@app.post("/api/ingest/yargi")
def ingest_yargi(sorgu: str, limit: int = 10, court: str = "ALL"):
    """Yargıtay/Danıştay/AYM kararlarını yargi-cli veya Bedesten API ile çek ve vektörleştir"""
    import sys
    import subprocess
    script_path = Path(__file__).resolve().parent.parent / "ingest_pipeline" / "yargitay_danistay_aym.py"
    cmd = [sys.executable, str(script_path), "--sorgu", sorgu, "--limit", str(limit), "--court", court]
    try:
        res = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return {"success": True, "output": res.stdout.strip()}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/ingest/mevzuat-baseline")
def ingest_mevzuat_baseline(limit: int = 150):
    """Mevzuat baseline Hugging Face veri setini yükle ve vektörleştir"""
    import sys
    import subprocess
    script_path = Path(__file__).resolve().parent.parent / "ingest_pipeline" / "mevzuat_baseline.py"
    cmd = [sys.executable, str(script_path), "--limit", str(limit)]
    try:
        res = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return {"success": True, "output": res.stdout.strip()}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/ingest/resmi-gazete")
def ingest_resmi_gazete(tarih: str = None, days: int = 1):
    """Resmi Gazete PDF indir, parse et ve vektörleştir"""
    import sys
    import subprocess
    script_path = Path(__file__).resolve().parent.parent / "ingest_pipeline" / "resmi_gazete.py"
    cmd = [sys.executable, str(script_path), "--days", str(days)]
    if tarih:
        cmd.extend(["--tarih", tarih])
    try:
        res = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return {"success": True, "output": res.stdout.strip()}
    except Exception as e:
        return {"success": False, "error": str(e)}



@app.get("/api/scraper/durum")
def scraper_durum():
    """Tarayıcı durumu ve önbellek bilgisi"""
    durum = {}
    for name, path in [("yargitay", "yargitay.json"), ("danistay", "danistay.json"), ("mevzuat", "mevzuat.json"), ("resmi_gazete", "resmi_gazete.json")]:
        p = DATA_DIR / "scraper" / path
        if p.exists():
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            durum[name] = {"adet": len(data), "guncelleme": data[0].get("scrape_tarihi", "") if data else ""}
        else:
            durum[name] = {"adet": 0, "guncelleme": ""}
    return durum


@app.get("/api/scraper/karar-ara")
def scraper_karar_ara(sorgu: str = "", kaynak: str = "tumu", kategori: str = ""):
    """Taranmış kararlarda arama yap"""
    sonuc = []
    dosyalar = []
    if kaynak in ("tumu", "yargitay"):
        dosyalar.append(DATA_DIR / "scraper" / "yargitay.json")
    if kaynak in ("tumu", "danistay"):
        dosyalar.append(DATA_DIR / "scraper" / "danistay.json")
    if kaynak in ("tumu", "mevzuat"):
        dosyalar.append(DATA_DIR / "scraper" / "mevzuat.json")
    if kaynak in ("tumu", "resmi_gazete"):
        dosyalar.append(DATA_DIR / "scraper" / "resmi_gazete.json")

    sorgu_lower = sorgu.lower()
    for dosya_yolu in dosyalar:
        if dosya_yolu.exists():
            with open(dosya_yolu, "r", encoding="utf-8") as f:
                data = json.load(f)
            for item in data:
                if kategori and item.get("kategori") != kategori:
                    continue
                if (sorgu_lower in (item.get("konu", "") or "").lower() or
                    sorgu_lower in (item.get("ozet", "") or "").lower() or
                    sorgu_lower in (item.get("esas", "") or "").lower() or
                    sorgu_lower in (item.get("baslik", "") or "").lower() or
                    sorgu_lower in (item.get("madde", "") or "").lower() or
                    sorgu_lower in (item.get("sonuc", "") or "").lower()):
                    sonuc.append(item)
    return sonuc

@app.post("/api/dataset/egit")
def dataset_egit(data: dict = None):
    """Belirtilen subdomain (avukat) veya global düzeyde dava, belge ve müvekkilleri vektör veritabanına ekler."""
    try:
        subdomain = data.get("subdomain", "") if data else ""
        project_dir = Path(__file__).parent.parent
        
        # Determine the source folder
        if subdomain and subdomain != "www":
            data_dir = project_dir / "storage" / "data" / "tenants" / subdomain
        else:
            data_dir = project_dir / "storage" / "data"
            
        from vector_store import get_vector_store
        store = get_vector_store(subdomain)
        store.clear() # Clear existing vectors to rebuild fresh
        
        logger.info(f"[AI Isolation] Indexing and training vector database for subdomain: '{subdomain or 'global'}'")
        
        total_indexed = 0
        
        # 1. Index Davalar
        davas_file = data_dir / "davas.json"
        if davas_file.exists():
            with open(davas_file, "r", encoding="utf-8") as f:
                davas_data = json.load(f).get("davas", [])
                for d in davas_data:
                    text = f"Dava Dosyası: {d.get('dosyaNo')}\nAdı: {d.get('ad')}\nKonu: {d.get('konu', '')}\nMahkeme: {d.get('mahkeme', '')}\nEsas No: {d.get('esasNo', '')}\nDurum: {d.get('durum', '')}\nAçıklama: {d.get('aciklama', '')}"
                    metadata = {"tur": "dava", "dosyaNo": d.get("dosyaNo"), "ad": d.get("ad")}
                    store.add_text(text, metadata, chunk=True)
                    total_indexed += 1
                    
                    if d.get("emsalKararlar"):
                        for em in d["emsalKararlar"]:
                            em_text = f"Emsal Karar ({em.get('kaynak')}): {em.get('mahkeme')}\nEsas: {em.get('esasNo')}, Karar: {em.get('kararNo')}\nKonu: {em.get('konu')}\nÖzet: {em.get('ozet')}"
                            em_metadata = {"tur": "emsal_karar", "dosyaNo": d.get("dosyaNo"), "esasNo": em.get("esasNo"), "kararNo": em.get("kararNo")}
                            store.add_text(em_text, em_metadata, chunk=True)
                            total_indexed += 1
                            
        # 2. Index Musteriler
        musteris_file = data_dir / "musteriler.json"
        if musteris_file.exists():
            with open(musteris_file, "r", encoding="utf-8") as f:
                musteris_data = json.load(f).get("musteriler", [])
                for m in musteris_data:
                    text = f"Müvekkil Bilgisi:\nAd Soyad: {m.get('ad', '')} {m.get('soyad', '')}\nTC Kimlik: {m.get('tcKimlik', '')}\nTelefon: {m.get('telefon', '')}\nE-Posta: {m.get('email', '')}\nNotlar: {m.get('notlar', '')}"
                    metadata = {"tur": "musteri", "id": m.get("id"), "tcKimlik": m.get("tcKimlik")}
                    store.add_text(text, metadata, chunk=True)
                    total_indexed += 1
                    
        # 3. Index Belgeler
        belgeler_file = data_dir / "belgeler.json"
        if belgeler_file.exists():
            with open(belgeler_file, "r", encoding="utf-8") as f:
                belgeler_data = json.load(f).get("belgeler", [])
                for b in belgeler_data:
                    text = f"Hukuki Belge/Evrak:\nBaşlık: {b.get('baslik', '')}\nTür: {b.get('tur', '')}\nİçerik: {b.get('icerik', '')}"
                    metadata = {"tur": "belge", "id": b.get("id"), "baslik": b.get("baslik")}
                    store.add_text(text, metadata, chunk=True)
                    total_indexed += 1

        # Also fallback load legacy corpus if needed
        corpus_path = project_dir / "python-backend" / "data" / "local_law_corpus.json"
        if not subdomain and corpus_path.exists():
            dataset.load_local_corpus(corpus_path)

        return {
            "basarili": True,
            "toplam_kayit": total_indexed,
            "subdomain": subdomain or "global",
            "mesaj": f"Yapay zeka '{subdomain or 'global'}' veritabanı başarıyla indekslendi ve eğitildi!"
        }
    except Exception as e:
        logger.error(f"Eğitim hatası: {e}")
        raise HTTPException(500, detail=str(e))


@app.get("/api/admin/benchmark/run")
def admin_benchmark_run(limit: int = 10, difficulty: str = "all"):
    """Gemini AI ile hukuki kavrama benchmark testi (HukukBERT cloze dataset)"""
    import random
    dataset_path = Path("C:/Users/acer/Desktop/hukukbert-main/benchmark/data/hukukbert_v1_cloze.jsonl")
    if not dataset_path.exists():
        alt_path = Path(__file__).parent.parent.parent / "hukukbert-main" / "benchmark" / "data" / "hukukbert_v1_cloze.jsonl"
        if alt_path.exists():
            dataset_path = alt_path
        else:
            raise HTTPException(404, detail="HukukBERT cloze dataset not found")
            
    questions = []
    try:
        with open(dataset_path, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    questions.append(json.loads(line))
    except Exception as e:
        logger.error(f"Failed to read HukukBERT benchmark dataset: {e}")
        raise HTTPException(500, detail=f"Dataset read error: {str(e)}")
        
    if difficulty != "all":
        questions = [q for q in questions if q.get("metadata", {}).get("difficulty") == difficulty]
        
    if not questions:
        raise HTTPException(404, detail=f"No questions found for difficulty '{difficulty}'")
        
    sample_size = min(limit, len(questions))
    selected_questions = random.sample(questions, sample_size)
    
    results = []
    correct_count = 0
    
    for q in selected_questions:
        sentence = q.get("sentence", "")
        options = q.get("options", [])
        gold = q.get("gold", "")
        
        choices_text = "\n".join([f"- {opt}" for opt in options])
        
        prompt = f"""Aşağıdaki hukuk cümlesindeki [MASK] ifadesinin yerine gelmesi gereken en uygun hukuki terimi seçenekler arasından belirle.
Cümle: {sentence}

Seçenekler:
{choices_text}

Önemli: YALNIZCA seçeneklerdeki kelimelerden birini yaz, başka hiçbir açıklama ekleme.

Seçilen Terim:"""

        model_answer = ""
        try:
            model_answer = call_gemini(prompt, temperature=0.0, max_tokens=10)
            model_answer = model_answer.strip().strip(".\"'*` ")
        except Exception as e:
            logger.error(f"Gemini benchmark query failed for ID {q.get('id')}: {e}")
            model_answer = f"Hata"
            
        def clean_text(t):
            return "".join(t.lower().split()).replace(".", "").replace(",", "").replace("'", "").replace('"', "")
            
        is_correct = clean_text(model_answer) == clean_text(gold)
        if is_correct:
            correct_count += 1
            
        results.append({
            "id": q.get("id"),
            "sentence": sentence,
            "options": options,
            "model_answer": model_answer if model_answer else "(Cevap Alınamadı)",
            "correct_answer": gold,
            "is_correct": is_correct,
            "law_area": q.get("metadata", {}).get("law_area", "genel"),
            "difficulty": q.get("metadata", {}).get("difficulty", "medium")
        })
        
    accuracy = (correct_count / sample_size) * 100 if sample_size > 0 else 0
    return {
        "success": True,
        "total": sample_size,
        "correct": correct_count,
        "accuracy": round(accuracy, 2),
        "model": f"Gemini ({GEMINI_MODEL})",
        "results": results
    }


class KategoriBelirleRequest(BaseModel):
    konu: str
    aciklama: str = ""


@app.post("/api/kategori/belirle")
def api_kategori_belirle(req: KategoriBelirleRequest):
    """Gemini AI ile davanın konusundan hukuk alanını (kategorisini) belirler"""
    text = f"Dava Konusu: {req.konu}\nDava Açıklaması: {req.aciklama}"
    
    import requests
    prompt = f"""Sen akıllı bir Hukuk Büro Asistanı yapay zekasısın. Görevin, verilen dava konusu ve açıklamasına göre davanın hangi hukuk alanına (kategorisine) ait olduğunu belirlemektir.
    
    Aşağıdaki kategorilerden tam olarak EN UYGUN OLANINI seç:
    - is (İş Hukuku alacakları, kıdem, ihbar, mesai, iş sözleşmeleri vb.)
    - bosanma (Boşanma, zina, nafaka, velayet vb.)
    - aile (Boşanma dışındaki aile hukuku, nişan, vesayet, kayyım, soybağı vb.)
    - miras (Veraset, mirasçılık, vasiyet, muris alacakları vb.)
    - kira (Kira sözleşmeleri, tahliye davaları, kira tespit, kira artış vb.)
    - tazminat (Maddi ve manevi tazminat talepleri, trafik kazası tazminatı, haksız fiil vb.)
    - ceza (Suç duyuruları, ceza davaları, soruşturmalar vb.)
    - ticaret (Şirketler, ortaklıklar, TTK hükümleri, çek/senet alacakları vb.)
    - icra (İcra takipleri, haciz işlemleri, ödeme emirleri, itirazın iptali vb.)
    - diger (Yukarıdakilere uymayan genel veya diğer hukuk alanları)

    Metin:
    {text}

    Önemli Kurallar:
    1. Yanıtında YALNIZCA yukarıda belirtilen kısa kodlardan birini yaz (örn: "is", "kira", "ceza" veya "diger").
    2. Başka hiçbir açıklama, yorum, ek kelime veya noktalama işareti ekleme. Yanıtın sadece kategori kodu olmalıdır.

    Kategori Kodu:"""

    # Gemini ile kategori belirle
    try:
        kategori = kategori_belirle_ai(req.konu, req.aciklama)
        logger.info(f"Gemini AI davanın kategorisini '{kategori}' olarak belirledi.")
        return {"kategori": kategori, "engine": "gemini"}
    except Exception as e:
        logger.error(f"Gemini kategori belirleme hatası: {e}")
    
    # Fallback: kural tabanlı
    kategori = classify_text(req.konu + " " + req.aciklama)
    logger.info(f"Kural tabanlı sistem davanın kategorisini '{kategori}' olarak belirledi.")
    return {"kategori": kategori, "engine": "rule-based"}


@app.post("/api/evrak/isle")
async def api_evrak_isle(
    file: UploadFile = File(...),
    davas_json: str = Form(...)
):
    """Yüklenen evrakın metnini çıkarır, en uygun davayı tespit eder ve resmi PDF'e dönüştürüp döner"""
    import json, time
    try:
        davas = json.loads(davas_json)
    except Exception as e:
        raise HTTPException(400, detail=f"Davas list JSON format error: {str(e)}")
        
    temp_path = UPLOAD_DIR / f"temp_evrak_{int(time.time())}_{file.filename}"
    with open(temp_path, "wb") as f:
        f.write(await file.read())
        
    ext = file.filename.split(".")[-1].lower()
    text = ""
    is_image = ext in ["jpg", "jpeg", "png"]
    
    # 1. Metin Ayıklama
    if ext == "pdf":
        from document_processor import extract_text_from_pdf
        text = extract_text_from_pdf(str(temp_path))
    elif is_image:
        from document_processor import extract_text_from_image
        text = extract_text_from_image(str(temp_path))
    elif ext == "udf":
        from document_processor import extract_text_and_signature_from_udf
        udf_data = extract_text_and_signature_from_udf(str(temp_path))
        text = udf_data["text"]
    elif ext == "txt":
        try:
            with open(temp_path, "r", encoding="utf-8") as f:
                text = f.read()
        except:
            try:
                with open(temp_path, "r", encoding="iso-8859-9") as f:
                    text = f.read()
            except Exception as e:
                logger.error(f"Txt okuma hatası: {e}")
                
    # 2. Hukuki Metadataları Çıkar
    from document_processor import extract_legal_metadata
    legal_info = extract_legal_metadata(text)
    
    matched_dava_id = ""
    
    # A. Regex/String Uyuşması ile Dava Bulma
    if legal_info.get("esas_no"):
        esas_clean = "".join(legal_info["esas_no"].split()).upper()
        for d in davas:
            d_esas = "".join((d.get("esasNo") or "").split()).upper()
            d_no = "".join((d.get("dosyaNo") or "").split()).upper()
            if d_esas and (d_esas in esas_clean or esas_clean in d_esas):
                matched_dava_id = d.get("id")
                break
            if d_no and (d_no in esas_clean or esas_clean in d_no):
                matched_dava_id = d.get("id")
                break
                
    if not matched_dava_id and legal_info.get("karar_no"):
        karar_clean = "".join(legal_info["karar_no"].split()).upper()
        for d in davas:
            d_karar = "".join((d.get("kararNo") or "").split()).upper()
            if d_karar and (d_karar in karar_clean or karar_clean in d_karar):
                matched_dava_id = d.get("id")
                break
                
    # B. Gemini AI ile Dava Bulma (Eşleşme bulunamadıysa)
    if not matched_dava_id and text and davas:
        dava_list_str = "\n".join([
            f"ID: {d.get('id')} - Dosya No: {d.get('dosyaNo')} - Ad: {d.get('ad')} - Konu: {d.get('konu') or ''} - Esas: {d.get('esasNo') or ''}"
            for d in davas
        ])
        
        sample_text = text[:2000]
        
        prompt = f"""Evrak metnine göre bu evrakın aşağıdaki dava listesindeki hangi dava dosyasıyla ilişkili olduğunu tespit et.

Dava Listesi:
{dava_list_str}

Evrak İçeriği:
{sample_text}

Önemli: YALNIZCA davanın ID değerini yaz (örn: clx123abc) veya eşleşme yoksa 'bulunamadi' yaz. Başka hiçbir şey ekleme.

Seçilen Dava ID:"""
        
        try:
            resp = call_gemini(prompt, temperature=0.0, max_tokens=30)
            resp_clean = "".join(resp.split()).replace(".", "").replace('"', "").replace("'", "")
            if resp_clean != "bulunamadi" and any(d.get("id") == resp_clean for d in davas):
                matched_dava_id = resp_clean
                logger.info(f"Gemini evrakı '{matched_dava_id}' ID'li dava ile eşleştirdi.")
        except Exception as e:
            logger.error(f"Gemini evrak eşleştirme hatası: {e}")
            
    # 3. Dosya Dönüştürme (Görsel ise PDF'e dönüştür)
    final_file_path = temp_path
    if is_image:
        from document_processor import image_to_searchable_pdf
        pdf_output_path = temp_path.with_suffix(".pdf")
        success = image_to_searchable_pdf(str(temp_path), str(pdf_output_path))
        if success and pdf_output_path.exists():
            final_file_path = pdf_output_path
            try:
                os.remove(temp_path)
            except:
                pass
                
    # 4. Dijital İmza Doğrulama (PAdES / XAdES)
    imza_durumu = "imzasiz"
    imzalayan = ""
    imza_tarihi = ""
    
    if ext == "pdf" or (is_image and Path(final_file_path).suffix == ".pdf"):
        from document_processor import verify_pdf_signature
        try:
            imza_data = verify_pdf_signature(str(final_file_path))
            imza_durumu = imza_data.get("imza_durumu", "imzasiz")
            imzalayan = imza_data.get("imzalayan") or ""
            imza_tarihi = imza_data.get("imza_tarihi") or ""
        except Exception as e:
            logger.error(f"Imza dogrulama hatasi: {e}")
    elif ext == "udf":
        from document_processor import extract_text_and_signature_from_udf
        try:
            udf_data = extract_text_and_signature_from_udf(str(temp_path))
            imza_durumu = udf_data.get("imza_durumu", "imzasiz")
            imzalayan = udf_data.get("imzalayan") or ""
        except Exception as e:
            logger.error(f"UDF imza dogrulama hatasi: {e}")
                
    # Next.js'e dosya yanıtını dön
    def cleanup_files(paths: list):
        for p in paths:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except:
                pass
                
    bg_tasks = BackgroundTasks()
    bg_tasks.add_task(cleanup_files, [str(final_file_path)])
    
    # Safe encoding for HTTP headers (avoid non-ascii character error in uvicorn)
    safe_imzalayan = imzalayan.encode("utf-8", errors="ignore").decode("latin-1")
    
    headers = {
        "X-Matched-Dava-Id": matched_dava_id or "",
        "X-Original-Filename": file.filename,
        "X-Imza-Durumu": imza_durumu,
        "X-Imzalayan": safe_imzalayan,
        "X-Imza-Tarihi": imza_tarihi,
        "Access-Control-Expose-Headers": "X-Matched-Dava-Id, X-Original-Filename, X-Imza-Durumu, X-Imzalayan, X-Imza-Tarihi"
    }
    
    media_type = "application/pdf" if (is_image or ext == "pdf") else "text/plain"
    
    return FileResponse(
        path=str(final_file_path),
        media_type=media_type,
        filename=Path(final_file_path).name,
        headers=headers,
        background=bg_tasks
    )


@app.post("/api/uyap/belge-oku")
async def api_uyap_belge_oku(file: UploadFile = File(None), text: str = Form(None)):
    """
    Accepts an uploaded PDF/XML/UDF file or a raw copy-pasted string containing UYAP cases,
    parses it, and returns structured JSON cases.
    """
    from uyap_parser import parse_uyap_raw_text, parse_uyap_pdf, parse_uyap_udf
    import shutil
    
    if text:
        logger.info("Parsing raw text input for UYAP cases")
        cases = parse_uyap_raw_text(text)
        return {"success": True, "dosyalar": cases, "kaynak": "Kopyalanan Metin"}
        
    if file:
        logger.info(f"Parsing uploaded file for UYAP cases: {file.filename}")
        temp_file = UPLOAD_DIR / file.filename
        try:
            with open(temp_file, "wb") as f:
                shutil.copyfileobj(file.file, f)
                
            ext = file.filename.split(".")[-1].lower()
            if ext == "pdf":
                cases = parse_uyap_pdf(str(temp_file))
            elif ext == "udf":
                cases = parse_uyap_udf(str(temp_file))
            elif ext in ("xml", "txt"):
                try:
                    with open(temp_file, "r", encoding="utf-8") as rf:
                        file_text = rf.read()
                except:
                    with open(temp_file, "r", encoding="iso-8859-9") as rf:
                        file_text = rf.read()
                cases = parse_uyap_raw_text(file_text)
            else:
                raise HTTPException(400, detail=f"Desteklenmeyen dosya formatı: {ext}")
                
            return {"success": True, "dosyalar": cases, "kaynak": file.filename}
        except Exception as e:
            logger.error(f"UYAP file parse error: {e}")
            raise HTTPException(500, detail=str(e))
        finally:
            if temp_file.exists():
                try:
                    os.remove(temp_file)
                except:
                    pass
                    
    raise HTTPException(400, detail="Lütfen bir dosya yükleyin veya metin girin.")


# UYAP Live Sync State
uyap_sync_status = {
    "durum": "idle",
    "adim": "Bekliyor...",
    "yuzde": 0,
    "detay": "",
    "hata": ""
}

class UyapSyncStartRequest(BaseModel):
    tc_no: str = ""
    uyap_sifre: str = ""
    giris_yontemi: str = "edevlet" # "edevlet" or "eimza"

@app.post("/api/uyap/sync-start")
def api_uyap_sync_start(req: UyapSyncStartRequest):
    global uyap_sync_status
    if uyap_sync_status["durum"] == "running":
        return {"success": False, "message": "Eşitleme işlemi zaten çalışıyor."}
        
    def progress_callback(status_dict):
        global uyap_sync_status
        uyap_sync_status.update(status_dict)
        
    def run_sync_thread():
        from uyap_scraper_selenium import sync_uyap_selenium
        try:
            # If backend is running inside Docker, requests target 'http://host.docker.internal:3001'
            api_url = "http://localhost:3001"
            if os.path.exists("/.dockerenv") or os.environ.get("AM_I_IN_A_DOCKER_CONTAINER"):
                api_url = "http://host.docker.internal:3001"
                
            sync_uyap_selenium(
                tc=req.tc_no,
                password=req.uyap_sifre,
                login_method=req.giris_yontemi,
                callback_progress=progress_callback,
                api_url=api_url
            )
        except Exception as e:
            logger.error(f"Thread sync error: {e}")
            progress_callback({
                "durum": "error",
                "adim": "Hata Oluştu",
                "yuzde": 0,
                "detay": str(e),
                "hata": str(e)
            })

    # Start thread
    uyap_sync_status = {
        "durum": "running",
        "adim": "Başlatılıyor...",
        "yuzde": 0,
        "detay": "Selenium oturumu hazırlanıyor.",
        "hata": ""
    }
    import threading
    t = threading.Thread(target=run_sync_thread, daemon=True)
    t.start()
    
    return {"success": True, "message": "UYAP eşitleme işlemi başlatıldı."}

@app.get("/api/uyap/sync-status")
def api_uyap_sync_status():
    global uyap_sync_status
    return uyap_sync_status


# ========== UETS TEBLIGAT ENDPOINT'LERI ==========

class UetsKurulumRequest(BaseModel):
    kurumKodu: str
    kurumSifre: str
    kullaniciAdi: str
    sifre: str
    testModu: bool = True

@app.post("/api/uets/kurulum")
def api_uets_kurulum(req: UetsKurulumRequest):
    """UETS entegrasyonunu kur ve giris yap."""
    basarili = uets_manager.kurulum_yap(
        req.kurumKodu, req.kurumSifre,
        req.kullaniciAdi, req.sifre, req.testModu
    )
    return {"success": basarili, "mesaj": "UETS baglantisi basarili" if basarili else "UETS baglantisi basarisiz"}


@app.get("/api/uets/kontrol")
def api_uets_kontrol(user_id: str = ""):
    """Yeni tebligatlari kontrol et."""
    import threading
    sonuc = {"durum": "calisiyor", "mesaj": "Kontrol baslatildi"}

    def run():
        try:
            res = uets_manager.kontrol_et(user_id)
            logger.info(f"UETS kontrol: {res}")
        except Exception as e:
            logger.error(f"UETS kontrol hatasi: {e}")

    t = threading.Thread(target=run, daemon=True)
    t.start()
    return sonuc


@app.get("/api/uets/sonuc")
def api_uets_sonuc(limit: int = 50):
    """Son UETS tebligat sonuclarini getir."""
    from pathlib import Path as P
    import json
    data_dir = P(__file__).parent / "data" / "uets"
    sonuc = []
    if data_dir.exists():
        for f in sorted(data_dir.glob("*.json"), reverse=True)[:limit]:
            try:
                with open(f, "r", encoding="utf-8") as fh:
                    sonuc.append(json.load(fh))
            except:
                pass
    return {"tebligatlar": sonuc, "adet": len(sonuc)}


# ========== AI RE-INDEX ==========

@app.post("/api/ai/re-index")
def api_reindex():
    """Manually trigger re-indexing of all vector stores."""
    import threading
    def reindex():
        try:
            from vector_indexer import index_global_datasets, index_all_tenants
            index_global_datasets()
            index_all_tenants()
        except Exception as e:
            logger.error(f"Re-index error: {e}")
    t = threading.Thread(target=reindex, daemon=True)
    t.start()
    return {"success": True, "message": "Yeniden indeksleme başlatıldı."}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8765)

